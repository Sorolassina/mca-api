from docx import Document
import re

def replace_ignore_case(text: str, replacements: dict) -> str:
    for old, new in replacements.items():
        pattern = re.compile(re.escape(old), re.IGNORECASE)
        text = pattern.sub(new, text)
    return text

def modifier_contenu_word(file_path: str, replacements: dict):
    try:
        doc = Document(file_path)
        modified = False

        # 📄 Paragraphes
        for para in doc.paragraphs:
            for run in para.runs:
                new_text = replace_ignore_case(run.text, replacements)
                if run.text != new_text:
                    run.text = new_text
                    modified = True

        # 📊 Tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            new_text = replace_ignore_case(run.text, replacements)
                            if run.text != new_text:
                                run.text = new_text
                                modified = True

        # 🧾 En-têtes et pieds de page
        for section in doc.sections:
            # En-tête
            for para in section.header.paragraphs:
                for run in para.runs:
                    new_text = replace_ignore_case(run.text, replacements)
                    if run.text != new_text:
                        run.text = new_text
                        modified = True

            # Pied de page
            for para in section.footer.paragraphs:
                for run in para.runs:
                    new_text = replace_ignore_case(run.text, replacements)
                    if run.text != new_text:
                        run.text = new_text
                        modified = True

        # 💾 Sauvegarde
        if modified:
            doc.save(file_path)

    except Exception as e:
        print(f"❌ Erreur traitement Word {file_path} : {e}")
